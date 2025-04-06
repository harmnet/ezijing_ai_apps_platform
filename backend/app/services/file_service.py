#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文件处理服务
提供从不同格式文件中提取文本的功能:
- PDF文件
- Word文档(.doc, .docx)
- Excel表格(.xls, .xlsx)
- 文本文件(.txt, .csv)
"""

import os
import re
import logging
from app.services.llm_service import MODEL_CONFIG

logger = logging.getLogger(__name__)

def get_token_limit_for_model(model_name):
    """获取模型的token上下文限制"""
    model_info = MODEL_CONFIG.get(model_name, {})
    return model_info.get("token_limit", 8192)  # 默认8192

def extract_text_from_file(file_path):
    """
    从文件中提取文本
    根据文件扩展名自动选择合适的解析方法
    
    Args:
        file_path: 文件路径
    
    Returns:
        str: 提取的文本内容
    """
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        return ""
    
    # 获取文件扩展名
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        # 根据文件扩展名选择相应的解析方法
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return extract_text_from_word(file_path)
        elif ext in ['.xls', '.xlsx']:
            return extract_text_from_excel(file_path)
        elif ext in ['.txt', '.csv', '.md', '.json', '.js', '.py', '.html', '.css']:
            return extract_text_from_text_file(file_path)
        else:
            logger.warning(f"不支持的文件格式: {ext}")
            return f"不支持解析该文件格式: {ext}"
    except Exception as e:
        logger.error(f"提取文本时出错: {str(e)}")
        return f"解析文件时发生错误: {str(e)}"

def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本"""
    try:
        # 尝试导入PyPDF2库
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                # 提取每一页的文本
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
                
                return clean_text(text)
        except ImportError:
            # 如果PyPDF2不可用，尝试使用pdfplumber
            try:
                import pdfplumber
                
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n\n"
                    
                    return clean_text(text)
            except ImportError:
                return "需要安装PyPDF2或pdfplumber库才能解析PDF文件。请运行: pip install PyPDF2 或 pip install pdfplumber"
    except Exception as e:
        logger.error(f"解析PDF时出错: {str(e)}")
        return f"解析PDF文件时发生错误: {str(e)}"

def extract_text_from_word(file_path):
    """从Word文档中提取文本"""
    try:
        # 尝试导入docx库（用于.docx文件）
        try:
            if file_path.lower().endswith('.docx'):
                import docx
                
                doc = docx.Document(file_path)
                text = ""
                
                # 提取文档中的所有段落
                for para in doc.paragraphs:
                    text += para.text + "\n"
                
                # 提取表格中的文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                    text += "\n"
                
                return clean_text(text)
            # 对于旧版的.doc文件，尝试使用textract
            elif file_path.lower().endswith('.doc'):
                try:
                    import textract
                    text = textract.process(file_path).decode('utf-8')
                    return clean_text(text)
                except ImportError:
                    return "需要安装textract库才能解析.doc文件。请运行: pip install textract"
        except ImportError:
            return "需要安装python-docx库才能解析Word文档。请运行: pip install python-docx"
    except Exception as e:
        logger.error(f"解析Word文档时出错: {str(e)}")
        return f"解析Word文档时发生错误: {str(e)}"

def extract_text_from_excel(file_path):
    """从Excel表格中提取文本"""
    try:
        # 尝试导入pandas库
        try:
            import pandas as pd
            
            # 读取Excel文件中的所有工作表
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            # 处理每个工作表
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text += f"工作表: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return clean_text(text)
        except ImportError:
            return "需要安装pandas库才能解析Excel表格。请运行: pip install pandas openpyxl"
    except Exception as e:
        logger.error(f"解析Excel表格时出错: {str(e)}")
        return f"解析Excel表格时发生错误: {str(e)}"

def extract_text_from_text_file(file_path):
    """从文本文件中提取文本"""
    try:
        # 尝试以UTF-8编码读取文件
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # 如果UTF-8解码失败，尝试GBK编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except UnicodeDecodeError:
                # 如果GBK也失败，尝试Latin-1编码（可以解码任何字节）
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
    except Exception as e:
        logger.error(f"读取文本文件时出错: {str(e)}")
        return f"读取文本文件时发生错误: {str(e)}"

def clean_text(text):
    """清理提取的文本"""
    if not text:
        return ""
    
    # 移除多余的空白行
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # 移除多余的空格
    text = re.sub(r'\s+', ' ', text)
    
    # 移除Unicode控制字符
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    return text.strip()
